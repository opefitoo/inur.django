import docx
from django.http import HttpResponse
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from invoices.modelspackage.invoice import get_default_invoicing_details


def add_dotted_line(doc):
    p = doc.add_paragraph()
    p_format = p.paragraph_format

    # Create the border element
    bottom_border = OxmlElement('w:pBdr')
    border = OxmlElement('w:bottom')
    border.set(qn('w:val'), 'dotted')
    border.set(qn('w:sz'), '6')  # Border size (1/8 pt, so this is 0.75 pt)
    border.set(qn('w:space'), '1')  # Space between text and border
    border.set(qn('w:color'), '000000')  # Border color (black)

    bottom_border.append(border)

    # Apply the border to the paragraph
    p._element.get_or_add_pPr().append(bottom_border)


def add_page_number(paragraph):
    # Create a PAGE field
    fldChar1 = OxmlElement('w:fldChar')  # Create the <w:fldChar> element
    fldChar1.set(qn('w:fldCharType'), 'begin')  # Set field type to begin

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')  # Create the <w:fldChar> element
    fldChar2.set(qn('w:fldCharType'), 'end')  # Set field type to end

    paragraph._element.append(fldChar1)
    paragraph._element.append(instrText)
    paragraph._element.append(fldChar2)


def add_field(paragraph, field_type):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field_type
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)

    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "1"  # This is a placeholder, Word will replace it
    run._r.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar4)


def add_header_with_line(doc):
    section = doc.sections[0]
    header = section.header

    # Add a paragraph to the header
    paragraph = header.paragraphs[0]

    # Add logo on the left side of the header
    run = paragraph.add_run()
    run.add_picture('invoices/static/images/Logo_SUR_quadri_transparent_pour_copas.png', width=docx.shared.Inches(0.75))

    # Add details on the right side of the header
    run = paragraph.add_run(get_default_invoicing_details_string_for_header())
    run.font.size = Pt(8)  # Smaller font size
    run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color (RGB 128, 128, 128)
    paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Add another paragraph for the line separator
    p = header.add_paragraph()

    # Create the border element for the line
    bottom_border = OxmlElement('w:pBdr')
    border = OxmlElement('w:bottom')
    border.set(qn('w:val'), 'single')  # Solid line
    border.set(qn('w:sz'), '6')  # Border size (1/8 pt, so this is 0.75 pt)
    border.set(qn('w:space'), '1')  # Space between text and border
    border.set(qn('w:color'), '000000')  # Border color (black)

    bottom_border.append(border)

    # Apply the border to the paragraph
    p._element.get_or_add_pPr().append(bottom_border)

def get_default_invoicing_details_string_for_header():
    default_invoicing_details = get_default_invoicing_details()
    if not default_invoicing_details:
        return "SUR.LU SARL"

    return f"{default_invoicing_details.name} {default_invoicing_details.get_full_address()} - email:{default_invoicing_details.email_address}/tel:{default_invoicing_details.phone_number}"



def generate_transfer_document(patientAnamnesis):
    # Create a new Document
    doc = Document()

    add_header_with_line(doc)

    # Create the footer with a table
    section = doc.sections[0]
    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=docx.shared.Inches(6.0))
    table.cell(0, 0).width = docx.shared.Inches(5.5)
    table.cell(0, 1).width = docx.shared.Inches(0.5)

    # Add text to the left side of the footer
    left_cell = table.cell(0, 0).paragraphs[0]
    run = left_cell.add_run(
        'Fiche de transfert de %s %s (%s)' % (patientAnamnesis.patient.first_name, patientAnamnesis.patient.name,
                                              patientAnamnesis.updated_on.strftime('%d-%m-%Y %H:%M')))
    run.italic = True
    run.font.size = Pt(8)  # Smaller font size
    run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color (RGB 128, 128, 128)

    table.cell(0, 0).vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.BOTTOM

    # Add page numbers to the right side of the footer
    right_cell = table.cell(0, 1).paragraphs[0]
    right_cell.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.RIGHT
    run.font.size = Pt(8)  # Smaller font size
    run = right_cell.add_run('Page ')
    add_field(right_cell, 'PAGE')
    run = right_cell.add_run('/')
    add_field(right_cell, 'NUMPAGES')

    table.cell(0, 1).vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.BOTTOM

    # Add a line to separate the header from the content
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(docx.enum.text.WD_BREAK.LINE)

    # Add a title
    doc.add_heading('FICHE DE TRANSFERT',
                    level=0)

    # Add the base data section
    doc.add_heading('Données de base', level=1)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('Nom Prénom : ')
    run.bold = True
    paragraph.add_run('%s %s' % (patientAnamnesis.patient.name, patientAnamnesis.patient.first_name))
    doc.add_paragraph('Adresse: %s' % patientAnamnesis.patient.get_full_address_date_based())
    doc.add_paragraph('Matricule : %s' % patientAnamnesis.patient.code_sn)
    doc.add_paragraph('Tél: %s' % patientAnamnesis.patient.phone_number)
    doc.add_paragraph('Caisse(s) de maladie : %s' % patientAnamnesis.patient.get_caisse_maladie_display())
    # translate it to french using
    doc.add_paragraph('Etat civil : %s' % patientAnamnesis.get_civil_status_display())
    # if is_under_dependence_insurance is True display "OUI" else "NON"
    doc.add_paragraph(
        'Assurance dépendance : %s' % ('OUI' if patientAnamnesis.patient.is_under_dependence_insurance else 'NON'))
    doc.add_paragraph('Degré de dépendance : %s ' % patientAnamnesis.dependance_insurance_level)
    doc.add_paragraph('Nationalité : %s ' % patientAnamnesis.nationality)
    doc.add_paragraph('Langue(s) parlée(s): %s' % patientAnamnesis.spoken_languages)
    doc.add_paragraph('Langue(s) comprise(s) : %s' % patientAnamnesis.spoken_languages)
    doc.add_paragraph('Régime de protection juridique: %s ' % patientAnamnesis.legal_protection_regimes)

    # Add the contact persons section
    doc.add_heading('Personne(s) de contact', level=1)
    contact_string_list = patientAnamnesis.get_list_of_beautiful_string_for_contact_persons()
    for contact_string in contact_string_list:
        doc.add_paragraph(contact_string)

    add_dotted_line(doc)

    # Add doctor and medical sections
    doc.add_heading('Médecin(s) traitant(s) :', level=1)
    if patientAnamnesis.get_list_of_beautiful_string_for_main_assigned_physicians():
        for doctor in patientAnamnesis.get_list_of_beautiful_string_for_main_assigned_physicians():
            doc.add_paragraph(doctor)
    else:
        doc.add_paragraph('Aucun médecin traitant assigné')
    # Other doctors section
    if patientAnamnesis.get_list_of_beautiful_string_for_other_assigned_physicians():
        doc.add_heading('Autre(s) médecin(s)', level=1)
        for doctor in patientAnamnesis.get_list_of_beautiful_string_for_other_assigned_physicians():
            doc.add_paragraph(doctor)

    add_dotted_line(doc)

    # Add medical data
    doc.add_heading('Données médicales', level=1)
    doc.add_heading('Pathologies', level=2)
    doc.add_paragraph(patientAnamnesis.pathologies)

    add_dotted_line(doc)

    doc.add_heading('Antécédents', level=2)
    doc.add_paragraph(patientAnamnesis.medical_background)

    # Add treatment
    doc.add_heading('Traitement', level=2)
    doc.add_paragraph(patientAnamnesis.treatments)

    # Add allergies
    doc.add_heading('Allergies', level=2)
    doc.add_paragraph(patientAnamnesis.allergies)

    # Add last parameters
    doc.add_heading('Paramètres', level=1)
    last_parameters = patientAnamnesis.get_last_tension_and_temperature_parameters()
    # Create a list of strings for each parameter, only if the parameter is not None
    parameters_list = []
    if last_parameters:
        if last_parameters.systolic_blood_press is not None and last_parameters.systolic_blood_press > 0:
            parameters_list.append(f"Tension max: {last_parameters.systolic_blood_press} mmHg")
        if last_parameters.diastolic_blood_press is not None and last_parameters.diastolic_blood_press > 0:
            parameters_list.append(f"Tension min: {last_parameters.diastolic_blood_press} mmHg")
        if last_parameters.heart_pulse is not None and last_parameters.heart_pulse > 0:
            parameters_list.append(f"Pouls: {last_parameters.heart_pulse} bpm")
        if last_parameters.temperature is not None and last_parameters.temperature > 0:
            parameters_list.append(f"Température: {last_parameters.temperature} °C")
        if last_parameters.stools_parameter is not None:
            parameters_list.append(f"Selles: {'OUI' if last_parameters.stools_parameter else 'NON'}")
        if last_parameters.vas is not None:
            parameters_list.append(f"EVA: {last_parameters.vas}")
        if last_parameters.weight is not None:
            parameters_list.append(f"Poids: {last_parameters.weight} kg")
        if last_parameters.oximeter_saturation is not None:
            parameters_list.append(f"Saturation O2: {last_parameters.oximeter_saturation} %")
        if last_parameters.blood_glucose is not None:
            parameters_list.append(f"Glycémie: {last_parameters.blood_glucose} g/L")
        if last_parameters.general_remarks:
            parameters_list.append(f"Remarques générales: {last_parameters.general_remarks}")
        parameters_list.append(f"Dernière mise à jour: {last_parameters.updated_on.strftime('%d/%m/%Y %H:%M')}")

    # Add the parameters to the document
    doc.add_paragraph('\n'.join(parameters_list))
    add_dotted_line(doc)
    # Add Aides techniques section
    doc.add_heading('Aides techniques', level=1)
    doc.add_paragraph(patientAnamnesis.technical_help)
    if patientAnamnesis.other_technical_help:
        doc.add_paragraph(patientAnamnesis.other_technical_help)
    doc.add_heading('Autres Aides techniques', level=2)
    doc.add_paragraph(patientAnamnesis.other_technical_help)
    # Add additional parameters if they are not None
    additional_parameters_list = []
    if patientAnamnesis.dental_prosthesis is not None:
        additional_parameters_list.append(f"Prothèses dentaires: {patientAnamnesis.get_dental_prosthesis_display()}")
    if patientAnamnesis.hearing_aid is not None:
        additional_parameters_list.append(f"Appareil auditif: {patientAnamnesis.get_hearing_aid_display()}")
    if patientAnamnesis.glasses is not None:
        # if True display "OUI" else "NON"
        additional_parameters_list.append(f"Lunettes: {'OUI' if patientAnamnesis.glasses else 'NON'}")
    if patientAnamnesis.other_prosthesis is not None:
        additional_parameters_list.append(f"Autres: {patientAnamnesis.other_prosthesis}")

    # Add the additional parameters to the document
    if additional_parameters_list:
        doc.add_paragraph('\n'.join(additional_parameters_list))

    # Add mobilization section
    doc.add_heading('Mobilisation', level=1)
    doc.add_paragraph(patientAnamnesis.get_mobilization_display() + " // " + f"Description: {patientAnamnesis.mobilization_description}")

    # Add nutrition autonomy section
    doc.add_heading('Autonomie alimentaire', level=1)
    if patientAnamnesis.nutrition_autonomy:
        doc.add_paragraph(patientAnamnesis.get_nutrition_autonomy_display() + " // " + f"Régime: {patientAnamnesis.diet}")
    else:
        doc.add_paragraph(f"Régime: {patientAnamnesis.diet}")

    # Add Soins d'hygiène section
    doc.add_heading('Soins d\'hygiène', level=1)
    doc.add_paragraph(patientAnamnesis.hygiene_care_location + "// Jours de Douche: " + patientAnamnesis.shower_days + " // " + f"Lavage cheveux: {patientAnamnesis.hair_wash_days}")
    doc.add_paragraph(f"Autres détails ou remarques: {patientAnamnesis.hygiene_general_remarks}")
    ## add dotted line
    add_dotted_line(doc)
    # Add elimitation section
    doc.add_heading('Elimination', level=1)
    # if TRUE display "OUI" else "NON"
    doc.add_paragraph( f"Incontinence Urinaire: {'OUI' if patientAnamnesis.urinary_incontinence else 'NON'} // " + f"Incontinence fécale: {'OUI' if patientAnamnesis.faecal_incontinence else 'NON'}")
    doc.add_paragraph(f"Protections: {'OUI' if patientAnamnesis.protection else 'NON'} // " + f"Protection Pendant la journée: {patientAnamnesis.day_protection} // " + f"Protection Pendant la nuit: {patientAnamnesis.night_protection}")

    other_elimination_param_list = []
    if patientAnamnesis.urinary_catheter:
        other_elimination_param_list.append("Sonde urinaire: OUI")
    if patientAnamnesis.crystofix_catheter:
        other_elimination_param_list.append("Crystofix: OUI")
    if patientAnamnesis.elimination_addnl_details:
        other_elimination_param_list.append(patientAnamnesis.elimination_addnl_details)
    doc.add_paragraph('\n'.join(other_elimination_param_list))

    # Add section for habits
    doc.add_heading('Habitudes', level=1)
    habits = patientAnamnesis.get_biography_habits()
    for habit in habits:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(habit)
        paragraph.style = 'List Bullet'

    # Save the document to a temporary in-memory buffer
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="Fiche_Transfert_%s_%s_du_%s.docx"' % (
        patientAnamnesis.patient.first_name, patientAnamnesis.patient.name,
        patientAnamnesis.updated_on.strftime('%d-%m-%Y'))

    doc.save(response)
    return response
